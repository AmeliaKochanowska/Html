import os
import logging
import json
import time
import uuid
import mimetypes
import magic
import requests
import html2text
import zipfile # Do eksportu danych konta
import base64  # Do obsługi zawartości plików w formacie base64
import sys


from flask import Flask, request, jsonify, send_file, render_template, send_from_directory, session, abort, redirect, url_for
from flask_babel import Babel, get_locale # Dodano Babel
from flask_cors import CORS
from authlib.integrations.flask_client import OAuth
from googleapiclient.discovery import build

# import gspread # Odkomentuj, jeśli używasz Google Sheets
# from oauth2client.service_account import ServiceAccountCredentials # Odkomentuj, jeśli używasz Google Sheets
from docx import Document
from docx.shared import Pt, Inches # RGBColor nie jest używane, można usunąć

from io import BytesIO
from bs4 import BeautifulSoup
from datetime import timedelta, datetime # Do ustawień sesji i formatowania daty
from pathlib import Path
from modules.q_openai import q_OpenAI


# --- Konfiguracja Logowania ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - [%(funcName)s] %(message)s') # Dodano funcName
logger = logging.getLogger(__name__)


# --- Konfiguracja Aplikacji Flask ---
app = Flask(__name__)
CORS(app, supports_credentials=True) # Umożliwia CORS i obsługę ciasteczek/sesji

# --- Klucz Sekretny dla Sesji ---
# !!! WAŻNE: W środowisku produkcyjnym użyj silnego, losowego klucza !!!
# Można go wygenerować np. za pomocą: python -c 'import os; print(os.urandom(24))'
# Najlepiej przechowywać go w zmiennej środowiskowej
secret = os.environ.get('FLASK_SECRET_KEY')
if not secret:
    logger.error('FLASK_SECRET_KEY is required')
    raise RuntimeError('Missing FLASK_SECRET_KEY')
app.secret_key = secret
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax' # Lub 'Strict' dla większego bezpieczeństwa
app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Dodane dla bezpieczeństwa
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7) # Czas trwania sesji

# --- Konfiguracja Babel (Internacjonalizacja) ---
app.config['BABEL_DEFAULT_LOCALE'] = 'en'
app.config['SUPPORTED_LANGUAGES'] = ['en', 'pl'] # Definicja wspieranych języków
babel = Babel() # Utwórz instancję Babel bez aplikacji

def get_locale_from_session():
    """Pobiera ustawienie języka z sesji użytkownika."""
    user_lang = session.get('language')
    if user_lang in app.config['SUPPORTED_LANGUAGES']:
        logger.debug(f"Using language from session: {user_lang}")
        return user_lang
    # Jeśli nie ma w sesji, można dodać fallback do `request.accept_languages`
    logger.debug(f"Language not in session or not supported, falling back to default: {app.config['BABEL_DEFAULT_LOCALE']}")
    return app.config['BABEL_DEFAULT_LOCALE']

# Teraz możemy przypisać funkcję do selektora
babel.locale_selector_func = get_locale_from_session
babel.init_app(app) # Zainicjalizuj Babel dla aplikacji

oauth = OAuth(app)
auth0 = oauth.register(
    'auth0',
    client_id=os.environ.get('AUTH0_CLIENT_ID'),
    client_secret=os.environ.get('AUTH0_CLIENT_SECRET'),
    client_kwargs={'scope': 'openid profile email'},
    server_metadata_url=f"https://{os.environ.get('AUTH0_DOMAIN')}/.well-known/openid-configuration",
)


# --- Konfiguracja OpenAI ---
_OPENAI_CLIENT = None
client = None
OPENAI_INIT_ERROR = None
assistant_id = "asst_7VvkAuYY9wsOmrc64Y05Kpjr"
GAMMA_API_URL = os.environ.get('GAMMA_API_URL', 'https://api.gamma.app/generate')

# WYMUŚ UŻYCIE KLUCZA Z RENDER - DODAJ TUTAJ
os.environ.pop('Q_FLAG_OPENAI_API_KEY_USE_DEFAULT', None)

logger.info("=== SZCZEGÓŁOWE DEBUGOWANIE OPENAI ===")
logger.info(f"OPENAI_API_KEY in environment: {'OPENAI_API_KEY' in os.environ}")
if 'OPENAI_API_KEY' in os.environ:
    api_key_sample = os.environ['OPENAI_API_KEY']
    logger.info(f"API key starts with: {api_key_sample[:10]}...")
    logger.info(f"API key length: {len(api_key_sample)}")
    logger.info(f"API key format valid: {api_key_sample.startswith('sk-')}")
else:
    logger.error("OPENAI_API_KEY not found in environment variables!")
    logger.info(f"Available environment variables: {list(os.environ.keys())}")
logger.info("==========================================")

logger.info(f"=== DEBUG INFO - PRZED INICJALIZACJĄ ===")
logger.info(f"OPENAI_API_KEY present: {'OPENAI_API_KEY' in os.environ}")
logger.info(f"OPENAI_API_KEY length: {len(os.environ.get('OPENAI_API_KEY', ''))}")
logger.info(f"Q_FLAG_OPENAI_API_KEY_USE_DEFAULT: {os.environ.get('Q_FLAG_OPENAI_API_KEY_USE_DEFAULT', 'Not set')}")
logger.info(f"Current working directory: {os.getcwd()}")
logger.info(f"Python path: {sys.path}")
logger.info(f"===========================================")

try:
    from modules.q_openai import q_OpenAI
    # WYMUSZENIE KLUCZA BEZPOŚREDNIO Z RENDER
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise Exception("OPENAI_API_KEY not found in environment")
    
    logger.info(f"Attempting to initialize q_OpenAI with key: {api_key[:10]}...")
    _OPENAI_CLIENT = q_OpenAI(assistant_id, api_key=api_key)
    
    # BEZPIECZNE PRZYPISANIE client
    if _OPENAI_CLIENT and _OPENAI_CLIENT.is_valid:
        client = _OPENAI_CLIENT.hnd_
        logger.info("OpenAI client zainicjalizowany pomyślnie")
        
        # Test file capability
        if not _OPENAI_CLIENT.enable_file_capability():
            logger.error('Asystent nie ma uprawnień do pracy z plikami')
    else:
        logger.error("q_OpenAI initialization failed - client is not valid")
        client = None
        OPENAI_INIT_ERROR = "q_OpenAI client initialization failed"
        
except ImportError as e:
    OPENAI_INIT_ERROR = f"Failed to import q_OpenAI module: {str(e)}"
    logger.error(OPENAI_INIT_ERROR)
    client = None
except Exception as ex:
    OPENAI_INIT_ERROR = str(ex)
    logger.error(f'Błąd inicjalizacji OpenAI: {OPENAI_INIT_ERROR}')
    client = None
    
logger.info(f"=== DEBUG INFO - PO INICJALIZACJI ===")
logger.info(f"_OPENAI_CLIENT: {_OPENAI_CLIENT}")
logger.info(f"client: {client}")
logger.info(f"OPENAI_INIT_ERROR: {OPENAI_INIT_ERROR}")
logger.info(f"======================================")

# --- Konfiguracja Google Search API ---
GOOGLE_API_KEY = os.environ.get('GOOGLE_API_KEY')
GOOGLE_CSE_ID = os.environ.get('GOOGLE_CSE_ID')

def google_search(query, num_results=10):
    """Wykonuje wyszukiwanie przez Google Custom Search API."""
    try:
        if not GOOGLE_API_KEY or not GOOGLE_CSE_ID:
            logger.error("Brak konfiguracji Google API - sprawdź GOOGLE_API_KEY i GOOGLE_CSE_ID")
            return []
            
        service = build("customsearch", "v1", developerKey=GOOGLE_API_KEY)
        
        result = service.cse().list(
            q=query,
            cx=GOOGLE_CSE_ID,
            num=min(num_results, 10)  # Google API pozwala max 10 na raz
        ).execute()
        
        search_results = []
        if 'items' in result:
            for item in result['items']:
                search_results.append({
                    'title': item.get('title', 'No title'),
                    'url': item.get('link', ''),
                    'snippet': item.get('snippet', 'No description available')
                })
                
        logger.info(f"Google Search zwróciło {len(search_results)} wyników dla: {query}")
        return search_results
        
    except Exception as e:
        logger.error(f"Błąd Google Search API: {e}")
        return []


# --- Przechowywanie Danych w Pamięci (Uproszczone) ---
# W produkcji użyj bazy danych!
conversation_threads = {}  # Słownik: session_id -> thread_id
uploaded_files = {}        # Słownik: file_id -> file_info
general_chat_history = {}  # Słownik dla historii ogólnego chatu, nie przypisanego do kategorii

# --- Przechowywanie Danych Użytkownika (Tymczasowe) ---
# W produkcji użyj bazy danych!
user_profiles = {}  # Słownik: user_id -> profile_data

# Prosta hierarchia planów - wyższa liczba oznacza więcej możliwości
PLAN_LEVELS = {
    'Deckster Free': 0,
    'AI Deck Builder': 1,
    'AI Deck Pro': 2,
    'Human Intelligence Precision': 3
}

# --- Katalog Upload ---
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    try:
        os.makedirs(UPLOAD_FOLDER)
        logger.info(f"Utworzono katalog: {UPLOAD_FOLDER}")
    except OSError as e:
        logger.error(f"Nie można utworzyć katalogu {UPLOAD_FOLDER}: {e}")
        # Można rozważyć zatrzymanie aplikacji, jeśli upload jest krytyczny

# --- Konfiguracja Google Sheets (Opcjonalna - Odkomentuj jeśli potrzebne) ---
# def get_google_sheets_client():
#     # ... (kod jak poprzednio) ...
# def save_to_google_sheets(data):
#     # ... (kod jak poprzednio) ...

# --- Funkcje Pomocnicze ---

def get_current_user_data():
    """Pobiera dane zalogowanego użytkownika z sesji."""
    logger.debug(f"Session contents: {dict(session)}")
    
    if 'user_id' in session:
        user_data = {
            'user_id': session.get('user_id'),
            'name': session.get('user_name', 'User'),
            'email': session.get('user_email', 'No Email'),
            'plan': session.get('user_plan', 'Deckster Free'),
            'language': session.get('language', app.config['BABEL_DEFAULT_LOCALE']),
            'notificationsEnabled': session.get('notifications_enabled', True),
            'thread_id': session.get('thread_id')
        }
        logger.debug(f"Returning user data: {user_data}")
        return user_data
    
    logger.debug("No user_id in session")
    return None

def require_login(f):
    """Dekorator do ochrony endpointów wymagających zalogowania."""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            logger.warning(f"Dostęp nieautoryzowany do {request.endpoint}")
            return abort(401, description="User not logged in.")
        return f(*args, **kwargs)
    return decorated_function

def require_plan(min_plan):
    """Dekorator sprawdzający czy użytkownik ma wymagany plan."""
    from functools import wraps
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            user_plan = session.get('user_plan', 'Deckster Free')
            if PLAN_LEVELS.get(user_plan, 0) < PLAN_LEVELS.get(min_plan, 0):
                logger.warning(
                    f"Użytkownik {session.get('user_id')} z planem {user_plan} próbował uzyskać dostęp do funkcji wymagającej planu {min_plan}"
                )
                return jsonify(success=False, error='Upgrade required'), 403
            return f(*args, **kwargs)
        return wrapper
    return decorator

def is_valid_thread_id(thread_id):
    """Waliduje format thread_id po stronie backendu."""
    if not thread_id or not isinstance(thread_id, str):
        return False
    if not thread_id.startswith('thread_'):
        return False
    if len(thread_id) < 20 or len(thread_id) > 50:
        return False
    
    # OpenAI używa alfanumeryczne + myślniki i podkreślniki
    import re
    valid_pattern = re.compile(r'^thread_[a-zA-Z0-9_-]+$')
    return bool(valid_pattern.match(thread_id))

def create_new_thread():
    """Tworzy nowy wątek konwersacji w OpenAI."""
    try:
        if _OPENAI_CLIENT and _OPENAI_CLIENT.is_valid:
            thread_id = _OPENAI_CLIENT.thread_start()
            if thread_id:
                logger.info(f"Utworzono nowy wątek OpenAI: {thread_id}")
                return thread_id
            else:
                logger.error("q_OpenAI.thread_start() zwróciło None")
                return None
        else:
            logger.error("_OPENAI_CLIENT nie jest zainicjalizowany lub nieprawidłowy")
            return None
    except Exception as e:
        logger.error(f"Błąd tworzenia wątku OpenAI: {e}")
        return None

# --- Funkcje pomocnicze do obsługi File Q&A ---

def upload_file_to_assistant(file_path):
    """Przesyła plik do OpenAI i zwraca jego identyfikator."""
    if client is None or _OPENAI_CLIENT is None:
        logger.error("OpenAI client nie jest zainicjalizowany")
        return None
        
    try:
        if _OPENAI_CLIENT.upload_file(file_path):
            return _OPENAI_CLIENT.last_response.id
    except Exception as e:
        logger.error(f"Błąd przesyłania pliku: {e}")
        return None
    return None

    # try:
    #     with open(file_path, 'rb') as f:
    #         response = client.files.create(
    #             file=f,
    #             purpose="assistants"
    #         )
    #         logger.info(f"Przesłano plik do OpenAI, uzyskano File ID: {response.id}")
    #         return response.id
    # except Exception as e:
    #     logger.exception(f"Błąd podczas przesyłania pliku do OpenAI: {e}")
    #     return None
        
def ask_question_about_file(file_id, question, thread_id=None):
    """Zadaje pytanie o zawartość pliku i zwraca odpowiedź AI."""
    if client is None:
        logger.error("OpenAI client nie jest zainicjalizowany")
        return {"error": "Service temporarily unavailable"}
    
    flag_created = False
    if not _OPENAI_CLIENT.thread_check(thread_id):
        thread_id = _OPENAI_CLIENT.thread_start()
        if thread_id is None:
            return {"error": "Failed to create thread"}
        flag_created = True

    try:
        # Dodaj pytanie do wątku
        message = client.beta.threads.messages.create(
            thread_id=thread_id,
            role="user",
            content=question,
            file_ids=[file_id]  # Tutaj dodajemy plik do wiadomości
        )
        logger.info(f"Dodano pytanie i plik {file_id} do wątku {thread_id}")
        
        # Uruchom asystenta
        run = client.beta.threads.runs.create(
            thread_id=thread_id,
            assistant_id=assistant_id,
            instructions="Please analyze the attached file and answer the user's question in detail."
        )
        logger.info(f"Uruchomiono asystenta (Run ID: {run.id}) w wątku {thread_id}")
        
        # Monitoruj status uruchomienia
        run_status = "queued"
        max_attempts = 120
        attempts = 0
        
        while run_status in ["queued", "in_progress"] and attempts < max_attempts:
            time.sleep(0.5)
            run_info = client.beta.threads.runs.retrieve(
                thread_id=thread_id,
                run_id=run.id
            )
            run_status = run_info.status
            attempts += 1
            if attempts % 10 == 0:
                logger.debug(f"Sprawdzanie statusu uruchomienia {run.id} (próba {attempts}): {run_status}")
        
        if run_status == "completed":
            logger.info(f"Uruchomienie {run.id} zakończone pomyślnie.")
            # Pobierz odpowiedź asystenta
            messages = client.beta.threads.messages.list(
                thread_id=thread_id,
                order="desc"
            )
            
            # Znajdź najnowszą wiadomość od asystenta
            assistant_messages = [msg for msg in messages.data if msg.role == "assistant"]
            
            if assistant_messages:
                latest_message = assistant_messages[0]
                if latest_message.content and len(latest_message.content) > 0:
                    answer = latest_message.content[0].text.value
                    logger.info(f"Otrzymano odpowiedź asystenta w wątku {thread_id}")
                    return {"answer": answer, "thread_id": thread_id}
                else:
                    error_msg = f"Najnowsza wiadomość asystenta w wątku {thread_id} ma nieoczekiwaną strukturę"
                    logger.error(error_msg)
                    return {"error": error_msg}
            else:
                error_msg = f"Brak odpowiedzi od asystenta w wątku {thread_id} po zakończeniu uruchomienia."
                logger.warning(error_msg)
                return {"error": error_msg}
        elif run_status == "failed":
            run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
            error_message = run_info.last_error.message if run_info.last_error else "Unknown error"
            logger.error(f"Uruchomienie asystenta {run.id} nie powiodło się: {error_message}")
            return {"error": f"Asystent napotkał błąd: {error_message}"}
        else:
            error_msg = f"Uruchomienie asystenta {run.id} zakończyło się statusem '{run_status}' po {attempts} próbach."
            logger.error(error_msg)
            return {"error": f"Asystent nie ukończył odpowiedzi w wymaganym czasie (status: {run_status})"}
    except Exception as e:
        logger.exception(f"Błąd podczas przetwarzania pytania o plik: {e}")
        return {"error": f"Error processing the file question: {str(e)}"}
    finally:
        if flag_created:
            _OPENAI_CLIENT.thread_stop(thread_id)
            
def format_message_with_context(message, topic=None):
    """Formatuje wiadomość dla AI dodając kontekst kategorii."""
    # (Kod jak poprzednio)
    prompt_templates = {
        "Problem": f"Pomóż mi opisać problem, który rozwiązuje mój produkt/usługa. Wykorzystaj te informacje: {message}",
        "Solution": f"Pomóż mi opisać rozwiązanie, które oferuje mój produkt/usługa. Wykorzystaj te informacje: {message}",
        "Why Now": f"Pomóż mi wyjaśnić, dlaczego teraz jest właściwy moment na mój produkt/usługa. Wykorzystaj te informacje: {message}",
        "Market Size": f"Pomóż mi oszacować rozmiar rynku (TAM, SAM, SOM) dla mojego produktu/usługi. Wykorzystaj te informacje: {message}",
        "Competition": f"Pomóż mi przeanalizować konkurencję dla mojego produktu/usługi. Wykorzystaj te informacje: {message}",
        "Product": f"Pomóż mi opisać mój produkt, jego funkcje i cechy. Wykorzystaj te informacje: {message}",
        "Business Model": f"Pomóż mi opisać model biznesowy mojego produktu/usługi. Wykorzystaj te informacje: {message}",
        "Team Founders & Management": f"Pomóż mi przedstawić zespół założycielski i zarządzający. Wykorzystaj te informacje: {message}",
        "Financials": f"Pomóż mi przygotować informacje finansowe dla mojego pitch decku. Wykorzystaj te informacje: {message}",
        "The Deal": f"Pomóż mi przedstawić warunki inwestycji dla mojego pitch decku. Wykorzystaj te informacje: {message}",
        "Company Purpose": f"Pomóż mi zdefiniować cel mojej firmy w jednym zwięzłym zdaniu. Wykorzystaj te informacje: {message}",
    }

    if topic and topic in prompt_templates:
        return prompt_templates[topic]

    # Jeśli nie ma konkretnego szablonu dla danej kategorii
    if topic:
        return f"[Topic: {topic}] Przygotuj treść na temat {topic} do pitch decku, korzystając z tych informacji: {message}"

    # Jeśli nie ma tematu (ogólny chat), zwróć wiadomość bez formatowania
    return message

# --- Routing ---

@app.route('/static/<path:path>')
def send_static(path):
    return send_from_directory('static', path)

# Usunięto /static/img - powinno być obsługiwane przez powyższy endpoint

@app.route('/')
def home():
    """Serwuje główną stronę HTML."""
    logger.info("Wyświetlanie strony głównej (/)")
    dev_mode = request.args.get('dev_mode', False)
    # Tutaj można sprawdzić sesję i przekazać dane użytkownika do szablonu, jeśli zalogowany
    user_data = get_current_user_data()
    return render_template('index.html', dev_mode=dev_mode, user=user_data,
                           openai_error=OPENAI_INIT_ERROR)

# Endpoint do sprawdzania stanu sesji (nowy)
@app.route('/check-session', methods=['GET'])
def check_session():
   """Sprawdza, czy użytkownik jest zalogowany i zwraca jego dane."""
   try:
       # DODANE SZCZEGÓŁOWE LOGOWANIE
       logger.info(f"Session check - Session ID: {session.get('_id', 'no session id')}")
       logger.info(f"Session contents: {dict(session)}")
       logger.info(f"Session permanent: {session.permanent}")
       
       # Sprawdź czy sesja istnieje i ma podstawowe dane
       user_id = session.get('user_id')
       if not user_id:
           logger.info("check_session: no user_id in session")
           return jsonify({
               'logged_in': False,
               'success': True,
               'message': 'No active session - no user_id',
               'debug_session': dict(session)  # Tymczasowo dla debugowania
           }), 200
       
       # POPRAWKA: Załaduj zapisane dane profilu
       saved_profile = user_profiles.get(user_id, {})
       
       # Pobierz wszystkie dane użytkownika
       user_data = {
           'user_id': user_id,
           'name': saved_profile.get('name') or session.get('user_name', 'User'),
           'email': session.get('user_email', 'No Email'),
           'plan': saved_profile.get('plan') or session.get('user_plan', 'Deckster Free'),
           'language': saved_profile.get('language') or session.get('language', app.config['BABEL_DEFAULT_LOCALE']),
           'notificationsEnabled': saved_profile.get('notifications_enabled', session.get('notifications_enabled', True)),
           'thread_id': session.get('thread_id')
       }
       
       # KRYTYCZNA WALIDACJA - sprawdź czy wszystkie wymagane dane są obecne
       required_fields = ['user_id', 'name', 'email']
       missing_fields = [field for field in required_fields if not user_data.get(field)]
       
       if missing_fields:
           logger.error(f"Missing required session fields: {missing_fields}")
           # Wyczyść niepełną sesję
           session.clear()
           return jsonify({
               'logged_in': False,
               'success': True,
               'message': f'Incomplete session data, missing: {missing_fields}'
           }), 200
       
       # POPRAWKA: Zaktualizuj sesję DOPIERO po walidacji
       session['user_name'] = user_data['name']
       session['user_plan'] = user_data['plan'] 
       session['language'] = user_data['language']
       session['notifications_enabled'] = user_data['notificationsEnabled']
       session.modified = True
       
       logger.info(f"Sesja aktywna dla użytkownika: {user_data.get('user_id')}")
       return jsonify({
           'logged_in': True, 
           'user_data': user_data,
           'success': True
       }), 200
       
   except Exception as e:
       logger.error(f"Error in check_session: {e}", exc_info=True)
       return jsonify({
           'logged_in': False,
           'success': False,
           'error': str(e)
       }), 500
        
@app.route('/login')
def login_auth0():
    return auth0.authorize_redirect(redirect_uri=os.environ.get('AUTH0_CALLBACK_URL'))


@app.route('/signup')
def signup_auth0():
    return auth0.authorize_redirect(
        redirect_uri=os.environ.get('AUTH0_CALLBACK_URL'),
        screen_hint='signup'
    )



@app.route('/callback')
def auth0_callback():
    try:
        logger.info("Auth0 callback started")
        token = auth0.authorize_access_token()
        logger.info("Token received successfully")
        
        userinfo = token.get('userinfo') or auth0.parse_id_token(token)
        logger.info(f"User info received: {userinfo.get('email', 'no email')}")
        
        # KRYTYCZNE: Sprawdź czy userinfo zawiera wymagane dane
        if not userinfo or not userinfo.get('sub'):
            logger.error("No user info or sub claim in token")
            return redirect(url_for('home') + '?error=invalid_token')
        
        # Wyczyść sesję PRZED ustawieniem nowych danych
        session.clear()
        
        # Ustaw dane użytkownika w sesji - POPRAWIONE NAZWY KLUCZY
        session['user_id'] = userinfo.get('sub')
        session['user_name'] = userinfo.get('name', userinfo.get('nickname', 'User'))
        session['user_email'] = userinfo.get('email', '')
        session['user_plan'] = 'Deckster Free'
        session['language'] = (userinfo.get('locale', app.config['BABEL_DEFAULT_LOCALE']) or 'en')[:2]
        session['notifications_enabled'] = True
        
        # Utwórz nowy thread - DODANA OBSŁUGA BŁĘDÓW
        try:
            thread_id = create_new_thread()
            if thread_id:
                session['thread_id'] = thread_id
                logger.info(f"Created thread: {thread_id}")
            else:
                logger.warning("Failed to create thread")
                # NIE TWÓRZ FALLBACK - zostaw puste, utworzy się przy pierwszym zapytaniu
                session['thread_id'] = None
        except Exception as e:
            logger.error(f"Error creating thread: {e}")
            session['thread_id'] = None
        
        # Ustaw sesję jako stałą
        session.permanent = True
        session.modified = True
        
        # KRYTYCZNE: Sprawdź czy sesja została zapisana
        if not session.get('user_id'):
            logger.error("CRITICAL: Session was not saved properly!")
            return redirect(url_for('home') + '?error=session_failed')
        
        logger.info(f"User {session['user_id']} logged in successfully")
        logger.info(f"Session after save: {dict(session)}")
        logger.info(f"Session ID: {session.get('_id', 'no session id')}")

        
        # POPRAWIONE PRZEKIEROWANIE - używaj redirect z parametrami
        return redirect(f"{url_for('home')}?login_success=1")
        
    except Exception as e:
        logger.error(f"Auth0 callback error: {e}", exc_info=True)
        # Wyczyść sesję w przypadku błędu
        session.clear()
        return redirect(f"{url_for('home')}?error=auth_failed")

        
# Endpoint do lokalnego logowania (placeholder)
@app.route('/login-local', methods=['POST'])
def login_local():
    # TODO: Zaimplementuj logikę logowania (sprawdzenie hasła, itp.)
    # Na razie symulujemy sukces i ustawiamy sesję
    data = request.json
    email = data.get('email')
    password = data.get('password') # W realnej aplikacji haszuj i porównuj!

    if email and password: # Bardzo prosta walidacja
        session.clear() # Wyczyść starą sesję
        session['user_id'] = f"user_{email.split('@')[0]}" # Proste ID
        session['user_name'] = "Logged In User" # Pobierz z bazy danych
        session['user_email'] = email
        session['user_plan'] = "Deckster Free" # Pobierz z bazy danych
        session['language'] = 'en' # Domyślny język po zalogowaniu
        session['notifications_enabled'] = True
        session['thread_id'] = create_new_thread() # Stwórz nowy wątek dla nowej sesji
        session.permanent = True # Ustaw sesję jako stałą (używa PERMANENT_SESSION_LIFETIME)
        logger.info(f"Użytkownik {email} zalogowany pomyślnie (symulacja). Thread ID: {session['thread_id']}")
        user_data = get_current_user_data()
        return jsonify(success=True, user_data=user_data)
    else:
        logger.warning(f"Nieudana próba logowania dla email: {email}")
        return jsonify(success=False, error="Invalid credentials"), 401

# Endpoint do wylogowania (nowy)
@app.route('/logout', methods=['GET', 'POST'])
@require_login
def logout():
    user_id = session.get('user_id')
    
    # Wyczyść sesję Flask kompletnie
    session.clear()
    session.modified = True
    
    # Przygotuj URL do Auth0 logout
    domain = os.environ.get('AUTH0_DOMAIN')
    client_id = os.environ.get('AUTH0_CLIENT_ID')
    return_to = url_for('home', _external=True)
    auth0_logout_url = f"https://{domain}/v2/logout?returnTo={return_to}&client_id={client_id}"
    
    logger.info(f"Użytkownik {user_id} wylogowany, przekierowanie do Auth0 logout")
    
    if request.method == 'POST':
        # Zwróć Auth0 logout URL do frontendu
        return jsonify({
            'success': True,
            'auth0_logout_url': auth0_logout_url,
            'message': 'Logged out successfully'
        })
    else:
        return redirect(auth0_logout_url)

def ensure_valid_thread(session_id, current_thread_id=None):
    """
    Zapewnia że mamy prawidłowy thread_id.
    Zwraca (thread_id, is_new_thread)
    """
    # Sprawdź czy podany thread_id jest prawidłowy
    if current_thread_id:
        if _OPENAI_CLIENT and _OPENAI_CLIENT.thread_check(current_thread_id):
            logger.debug(f"Thread {current_thread_id} is valid and exists")
            return current_thread_id, False
        else:
            logger.warning(f"Thread {current_thread_id} has valid format but doesn't exist in OpenAI")
    elif current_thread_id:
        logger.warning(f"Thread {current_thread_id} has invalid format")
    
    # Jeśli nie mamy prawidłowego thread_id, utwórz nowy
    new_thread_id = create_new_thread()
    if not new_thread_id:
        logger.error(f"Failed to create new thread for session {session_id}")
        raise RuntimeError("Failed to create conversation thread")
    
    logger.info(f"Created new thread {new_thread_id} for session {session_id}")
    return new_thread_id, True


@app.route('/ask', methods=['POST'])
@require_login
def ask():
    """Obsługuje zapytania do asystenta AI z poprawną obsługą thread_id."""
    if client is None:
        logger.error("OpenAI client nie jest zainicjalizowany")
        return jsonify({"error": "Service temporarily unavailable"}), 503

    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        data = request.json
        question = data.get("question")
        topic = data.get("topic", "")
        request_thread_id = data.get("threadId")  # Thread z frontendu

        if not question:
            return jsonify({"error": "Brak pytania"}), 400

        # UJEDNOLICONA OBSŁUGA THREAD_ID
        current_thread_id = request_thread_id or session.get('thread_id')
        
        try:
            thread_id, is_new = ensure_valid_thread(session_id, current_thread_id)
            
            # Aktualizuj sesję jeśli thread się zmienił
            if thread_id != session.get('thread_id'):
                session['thread_id'] = thread_id
                session.modified = True
                logger.info(f"Updated session thread_id to {thread_id}")
                
        except RuntimeError as e:
            logger.error(f"Thread management failed for {session_id}: {e}")
            return jsonify({"error": "AI Assistant currently unavailable"}), 503

        # Formatuj wiadomość według kontekstu
        formatted_question = format_message_with_context(question, topic)

        # Zapisz w historii
        if topic:
            if 'conversation_history' not in session:
                session['conversation_history'] = {}
            if topic not in session['conversation_history']:
                session['conversation_history'][topic] = []
            if not isinstance(session['conversation_history'][topic], list):
                session['conversation_history'][topic] = []
            session['conversation_history'][topic].append({
                "question": question, 
                "timestamp": time.time()
            })
            session.modified = True
        else:
            user_id = session_id
            if user_id not in general_chat_history:
                general_chat_history[user_id] = []
            general_chat_history[user_id].append({
                "question": question, 
                "timestamp": time.time()
            })

        # Wyślij do OpenAI
        try:
            client.beta.threads.messages.create(
                thread_id=thread_id,
                role="user",
                content=formatted_question
            )

            run = client.beta.threads.runs.create(
                thread_id=thread_id,
                assistant_id=assistant_id
            )

            # Monitoruj wykonanie
            run_status = "queued"
            max_attempts = 120
            attempts = 0

            while run_status in ["queued", "in_progress"] and attempts < max_attempts:
                time.sleep(0.5)
                run_info = client.beta.threads.runs.retrieve(
                    thread_id=thread_id,
                    run_id=run.id
                )
                run_status = run_info.status
                attempts += 1

            if run_status == "completed":
                messages = client.beta.threads.messages.list(
                    thread_id=thread_id,
                    order="desc"
                )

                assistant_messages = [msg for msg in messages.data if msg.role == "assistant"]

                if assistant_messages and assistant_messages[0].content:
                    latest_message = assistant_messages[0]
                    if hasattr(latest_message.content[0], 'text'):
                        answer = latest_message.content[0].text.value

                        # Zapisz odpowiedź w historii
                        if topic and topic in session['conversation_history']:
                            latest_entry = session['conversation_history'][topic][-1]
                            latest_entry["answer"] = answer
                            session.modified = True
                        elif user_id in general_chat_history:
                            general_chat_history[user_id][-1]["answer"] = answer

                        # ZAWSZE zwróć thread_id w odpowiedzi
                        return jsonify({
                            "answer": answer, 
                            "threadId": thread_id
                        })
                    else:
                        logger.error(f"Unexpected message structure from assistant")
                        return jsonify({"error": "Asystent zwrócił nieprawidłowy format odpowiedzi"}), 500
                else:
                    return jsonify({"error": "Brak odpowiedzi od asystenta"}), 500
                    
            elif run_status == "failed":
                run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
                error_message = run_info.last_error.message if run_info.last_error else "Unknown error"
                return jsonify({"error": f"Asystent napotkał błąd: {error_message}"}), 500
            else:
                return jsonify({"error": f"Timeout lub nieoczekiwany status: {run_status}"}), 504

        except Exception as e:
            logger.exception(f"OpenAI API error for {session_id}: {e}")
            return jsonify({"error": f"Błąd komunikacji z asystentem: {str(e)}"}), 500

    except Exception as e:
        logger.exception(f"General error in /ask for {session_id}: {e}")
        return jsonify({"error": f"Błąd serwera: {str(e)}"}), 500

@app.route('/upload-file', methods=['POST'])
@require_login
def upload_file_endpoint(): # Zmieniono nazwę funkcji, aby uniknąć konfliktu
    """Obsługuje przesyłanie plików przez użytkownika."""
    logger.info(f"--start upload--")

    user_data = get_current_user_data()
    logger.info(f"user_data = {user_data}")

    session_id = user_data['user_id']
    thread_id = session.get('thread_id', 'unknown_thread') # Użyj wątku z sesji

    try:
        if 'file' not in request.files:
            logger.warning(f"Brak pliku w zapytaniu od {session_id}")
            return jsonify({"success": False, "error": "No file part"}), 400

        file = request.files['file']
        if file.filename == '':
            logger.warning(f"Przesłano plik bez nazwy od {session_id}")
            return jsonify({"success": False, "error": "No selected file"}), 400

        category = request.form.get('category', 'general')
        logger.info(f"Otrzymano plik '{file.filename}' od {session_id} dla kategorii '{category}' (wątek: {thread_id})")

        # Tworzymy unikalny identyfikator pliku
        file_id = str(uuid.uuid4())

        # Tworzymy katalog dla użytkownika, jeśli nie istnieje
        user_folder = os.path.join(UPLOAD_FOLDER, session_id)
        if not os.path.exists(user_folder):
            os.makedirs(user_folder)

        # Ścieżka do pliku
        # Używamy file_id jako nazwy, aby uniknąć kolizji i problemów ze znakami specjalnymi
        file_extension = os.path.splitext(file.filename)[1]
        safe_filename = file_id + file_extension
        file_path = os.path.join(user_folder, safe_filename)

        # Zapisujemy plik
        file.save(file_path)
        logger.info(f"Zapisano plik jako: {file_path}")

        # Zapisujemy informacje o pliku (w pamięci - użyj bazy danych w produkcji)
        file_info = {
            'original_name': file.filename,
            'path': file_path,
            'user_id': session_id, # Powiąż z użytkownikiem
            'thread_id': thread_id, # Powiąż z wątkiem (opcjonalnie)
            'category': category,
            'mime_type': file.content_type or mimetypes.guess_type(file.filename)[0] or 'application/octet-stream',
            'size': os.path.getsize(file_path),
            'added': time.time()
        }
        logger.info(f"file_info = {file_info}")

        uploaded_files[file_id] = file_info
        logger.debug(f"Zapisano metadane pliku {file_id}: {file_info}")

        # Utwórz kapsułkę danych pliku do użytku przez frontend
        file_data = {}
        
        # Odczytaj zawartość pliku w zależności od typu
        try:
            if file_info['mime_type'].startswith('text/'):
                # Dla plików tekstowych
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_data['content'] = f.read()
            elif file_info['mime_type'].startswith('image/'):
                # Dla obrazów - konwersja do base64
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                    file_data['content'] = base64.b64encode(file_content).decode('utf-8')
                    file_data['contentType'] = file_info['mime_type']
            # Dla innych typów plików można dodać więcej obsługi
        except Exception as e:
            logger.warning(f"Nie udało się odczytać zawartości pliku {file_id}: {e}")
            # Nie zwracaj błędu, po prostu pomiń zawartość pliku

        # Do odpowiedzi zwracamy tylko niesensytywne metadane
        ret = jsonify({
            "success": True,
            "fileId": file_id,
            "originalName": file.filename,
            "fileData": file_data
        })

        logger.info(f"ret = {ret}")

        return ret

    except Exception as e:
        logger.exception(f"Błąd podczas przesyłania pliku od {session_id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/download-file/<file_id>', methods=['GET'])
@require_login
def download_file_endpoint(file_id): # Zmieniono nazwę funkcji
    """Umożliwia pobranie wcześniej przesłanego pliku."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    logger.info(f"Żądanie pobrania pliku {file_id} przez użytkownika {session_id}")

    try:
        if file_id not in uploaded_files:
            logger.warning(f"Plik {file_id} nie znaleziony w uploaded_files.")
            abort(404, description="File not found in metadata.")

        file_info = uploaded_files[file_id]

        # Sprawdź, czy użytkownik ma prawo dostępu do tego pliku
        if file_info.get('user_id') != session_id:
            logger.warning(f"Użytkownik {session_id} próbuje pobrać plik ({file_id}) należący do {file_info.get('user_id')}")
            abort(403, description="Access denied to this file.") # Forbidden

        file_path = file_info.get('path')
        if not file_path or not os.path.exists(file_path):
             logger.error(f"Ścieżka do pliku {file_id} ({file_path}) nie istnieje lub jest niepoprawna.")
             # Usuń błędny wpis z metadanych
             if file_id in uploaded_files: del uploaded_files[file_id]
             abort(404, description="File path not found or invalid.")


        logger.info(f"Wysyłam plik: {file_path} jako {file_info['original_name']}")
        return send_file(
            file_path,
            as_attachment=True,
            download_name=file_info['original_name']
            # mimetype=file_info['mime_type'] # send_file zazwyczaj dobrze zgaduje mimetype
        )

    except Exception as e:
        logger.exception(f"Błąd podczas pobierania pliku {file_id}: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/view-file/<file_id>', methods=['GET'])
@require_login
def view_file_endpoint(file_id): # Zmieniono nazwę funkcji
    """Umożliwia podgląd pliku w przeglądarce, jeśli to możliwe."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    logger.info(f"Żądanie podglądu pliku {file_id} przez użytkownika {session_id}")

    try:
        if file_id not in uploaded_files:
             logger.warning(f"Plik {file_id} nie znaleziony w uploaded_files (view).")
             abort(404, description="File not found.")

        file_info = uploaded_files[file_id]

        if file_info.get('user_id') != session_id:
            logger.warning(f"Użytkownik {session_id} próbuje podejrzeć plik ({file_id}) należący do {file_info.get('user_id')}")
            abort(403, description="Access denied.")

        file_path = file_info.get('path')
        if not file_path or not os.path.exists(file_path):
             logger.error(f"Ścieżka do pliku {file_id} ({file_path}) nie istnieje lub jest niepoprawna (view).")
             if file_id in uploaded_files: del uploaded_files[file_id]
             abort(404, description="File path not found or invalid.")

        mime_type = file_info.get('mime_type', 'application/octet-stream')
        can_preview_inline = (
            mime_type.startswith(('image/', 'text/', 'audio/', 'video/')) or
            mime_type == 'application/pdf'
            )

        if can_preview_inline:
            logger.info(f"Wysyłam plik {file_path} do podglądu (inline, mimetype: {mime_type})")
            return send_file(
                file_path,
                mimetype=mime_type,
                as_attachment=False # Kluczowe dla podglądu inline
            )
        else:
            logger.info(f"Nie można podejrzeć pliku {file_path} (mimetype: {mime_type}), wysyłam jako załącznik.")
            # Jeśli nie można podejrzeć, pobieramy plik
            return send_file(
                file_path,
                as_attachment=True,
                download_name=file_info['original_name']
            )

    except Exception as e:
        logger.exception(f"Błąd podczas podglądu pliku {file_id}: {e}")
        return jsonify({"error": str(e)}), 500
        

@app.route('/process-files', methods=['POST'])
@require_login
def process_files_endpoint():
    """Przetwarza listę przesłanych plików przez AI."""
    if client is None:
        logger.error("OpenAI client nie jest zainicjalizowany")
        return jsonify({"error": "Service temporarily unavailable"}), 503
    
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    thread_id = session.get('thread_id')

    if not is_valid_thread_id(thread_id) or not _OPENAI_CLIENT.thread_check(thread_id):
        logger.warning(f"Nieprawidłowy lub nieistniejący thread_id: {thread_id}. Tworzę nowy.")
        thread_id = create_new_thread()
        if not thread_id:
            return jsonify({"error": "AI Assistant currently unavailable"}), 503
        session['thread_id'] = thread_id
        session.modified = True
        logger.info(f"Utworzono nowy thread_id dla process_files: {thread_id}")


    try:
        data = request.json
        file_infos_from_client = data.get('files', [])
        category = data.get('category')
        message = data.get('message', '')

        if not file_infos_from_client:
            return jsonify({"error": "No file information provided"}), 400
        if not thread_id:
            logger.error(f"Brak thread_id w sesji dla {session_id} podczas przetwarzania plików.")
            return jsonify({"error": "No active conversation thread found."}), 400

        logger.info(f"Przetwarzanie {len(file_infos_from_client)} plików dla {session_id} w kat. '{category}' (wątek {thread_id})")

        # Lista OpenAI file_ids do przetworzenia
        openai_file_ids = []
        filenames = []
        
        # Przetwórz każdy plik
        for file_client_info in file_infos_from_client:
            file_id = file_client_info.get('fileId')
            if file_id in uploaded_files:
                file_meta = uploaded_files[file_id]
                if file_meta.get('user_id') == session_id:
                    file_path = file_meta.get('path')
                    
                    # Dodaj plik do OpenAI (tylko jeśli nie był wcześniej dodany)
                    if 'openai_file_id' not in file_meta:
                        openai_file_id = upload_file_to_assistant(file_path)
                        if openai_file_id:
                            # Zapisz ID pliku z OpenAI z powrotem do metadanych
                            file_meta['openai_file_id'] = openai_file_id
                            uploaded_files[file_id] = file_meta
                            openai_file_ids.append(openai_file_id)
                            filenames.append(file_meta.get('original_name', 'Unknown'))
                    else:
                        # Użyj już istniejącego ID pliku OpenAI
                        openai_file_id = file_meta['openai_file_id']
                        openai_file_ids.append(openai_file_id)
                        filenames.append(file_meta.get('original_name', 'Unknown'))
                else:
                    logger.warning(f"Użytkownik {session_id} próbuje przetworzyć plik {file_id}, do którego nie ma dostępu.")
            else:
                logger.warning(f"Plik o ID {file_client_info.get('fileId')} z zapytania nie znaleziony w metadanych.")

        if not openai_file_ids:
            return jsonify({"error": "No valid files found for processing."}), 400
            
        # Przygotuj pytanie
        if message:
            prompt = message
        else:
            prompt = f"Analyze these files: {', '.join(filenames)}"
            
        # Dodaj wiadomość do wątku z plikami
        try:
            client.beta.threads.messages.create(
                thread_id=thread_id,
                role="user",
                content=prompt,
                file_ids=openai_file_ids  # TUTAJ jest kluczowa zmiana - dołączamy pliki do wiadomości
            )
            logger.info(f"Dodano wiadomość z {len(openai_file_ids)} plikami do wątku {thread_id}")
            
            # Uruchom asystenta
            run = client.beta.threads.runs.create(
                thread_id=thread_id,
                assistant_id=assistant_id,
                instructions="Analyze the attached files thoroughly and provide a detailed response."
            )
            logger.info(f"Uruchomiono asystenta (Run ID: {run.id}) dla analizy plików")
            
            # Monitoruj status (jak wcześniej)
            run_status = "queued"
            max_attempts = 120
            attempts = 0
            
            while run_status in ["queued", "in_progress"] and attempts < max_attempts:
                time.sleep(0.5)
                run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
                run_status = run_info.status
                attempts += 1
                if attempts % 10 == 0:
                    logger.debug(f"Status analizy plików: {run_status} (próba {attempts})")
            
            if run_status == "completed":
                # Pobierz odpowiedź asystenta
                messages = client.beta.threads.messages.list(thread_id=thread_id, order="desc")
                assistant_messages = [msg for msg in messages.data if msg.role == "assistant"]
                
                if assistant_messages and assistant_messages[0].content:
                    answer = assistant_messages[0].content[0].text.value
                    logger.info(f"Otrzymano odpowiedź AI po analizie plików.")
                    
                    return jsonify({"success": True, "answer": answer})
                else:
                    return jsonify({"error": "No response from assistant"}), 500
            else:
                error_message = f"Assistant failed to complete (status: {run_status})"
                if run_status == "failed":
                    run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
                    if run_info.last_error:
                        error_message = run_info.last_error.message
                
                logger.error(f"File analysis failed: {error_message}")
                return jsonify({"error": error_message}), 500
                
        except Exception as e:
            logger.exception(f"Error processing files with OpenAI: {e}")
            return jsonify({"error": f"Error processing files: {str(e)}"}), 500
            
    except Exception as e:
        logger.exception(f"General error in process_files_endpoint: {e}")
        return jsonify({"error": str(e)}), 500
        
@app.route('/generate-docx', methods=['POST'])
@require_login
def generate_docx_endpoint(): # Zmieniono nazwę funkcji
    """Generuje plik DOCX na podstawie zapisanych odpowiedzi."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    logger.info(f"Żądanie wygenerowania DOCX przez użytkownika {session_id}")

    try:
        # Pobierz zapisane sekcje z frontendu
        data = request.json
        sections = data.get('sections', {}) # Oczekujemy struktury: {'Section Name': [{'content': '...', 'timestamp': '...'}, ...]}

        if not sections:
             return jsonify({"error": "No sections provided for DOCX generation."}), 400

        # Utwórz nowy dokument
        doc = Document()

        # Ustawienia dokumentu (opcjonalnie)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri' # Lub inny font
        font.size = Pt(11)

        # Nagłówek dokumentu
        title = doc.add_heading('Pitch Deck Content Draft', 0)
        title.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER (1)

        # Dodaj datę wygenerowania
        from datetime import datetime
        current_date = datetime.now().strftime("%Y-%m-%d %H:%M")
        date_paragraph = doc.add_paragraph(f"Generated by Deckster AI on: {current_date}")
        date_paragraph.alignment = 1
        doc.add_paragraph()  # Pusty wiersz dla odstępu

        # Dodaj sekcje z treścią
        # Sortuj sekcje alfabetycznie dla spójności
        for section_name in sorted(sections.keys()):
            content_list = sections[section_name]
            if not content_list: continue # Pomiń puste sekcje

            doc.add_heading(section_name, level=1) # Nagłówek sekcji

            # Dodaj treść dla każdej pozycji w sekcji
            for content_item in content_list:
                text = content_item.get('content', '').strip()
                if text: # Dodaj tylko niepustą treść
                    doc.add_paragraph(text)
                    # Dodaj mały odstęp między elementami w sekcji (opcjonalnie)
                    # doc.add_paragraph().add_run().font.size = Pt(6)

            doc.add_paragraph() # Odstęp po całej sekcji

        # Zapisz dokument do pamięci (BytesIO)
        mem_file = BytesIO()
        doc.save(mem_file)
        mem_file.seek(0) # Przewiń na początek pliku w pamięci

        logger.info(f"Wygenerowano plik DOCX dla użytkownika {session_id}")
        return send_file(
            mem_file,
            as_attachment=True,
            download_name=f"deckster_ai_draft_{datetime.now().strftime('%Y%m%d')}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        logger.exception(f"Błąd podczas generowania pliku DOCX dla {session_id}: {e}")
        return jsonify({"error": f"Error generating DOCX file: {str(e)}"}), 500

@app.route('/generate-presentation', methods=['POST'])
@require_login
def generate_presentation():
    """Generuje prezentację w zewnętrznym systemie (np. Gamma)."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    user_plan = session.get('user_plan')

    if user_plan != 'AI Deck Pro':
        logger.warning(f"Użytkownik {session_id} z planem '{user_plan}' próbował wygenerować prezentację")
        return jsonify(success=False, error="Presentation generation is available only for AI Deck Pro plan."), 403

    try:
        data = request.json
        sections = data.get('sections', {})
        if not sections:
            return jsonify(success=False, error="No sections provided"), 400

        payload = {
            'user_id': session_id,
            'sections': sections
        }
        logger.info(f"Wysyłanie danych do Gamma API dla użytkownika {session_id}")

        response = requests.post(GAMMA_API_URL, json=payload, timeout=30)
        response.raise_for_status()
        result = response.json()

        link = result.get('link') or result.get('url')
        if not link:
            logger.error(f"Gamma API nie zwróciło linku dla {session_id}: {result}")
            return jsonify(success=False, error="No presentation link returned"), 500

        logger.info(f"Wygenerowano prezentację dla {session_id}: {link}")
        return jsonify(success=True, link=link)

    except Exception as e:
        logger.exception(f"Błąd podczas generowania prezentacji dla {session_id}: {e}")
        return jsonify(success=False, error=str(e)), 500

@app.route('/list-user-files', methods=['GET'])
@require_login
def list_user_files_endpoint(): # Zmieniono nazwę funkcji
    """Zwraca listę plików przesłanych przez zalogowanego użytkownika."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    logger.info(f"Żądanie listy plików przez użytkownika {session_id}")

    try:
        user_files_list = []
        # Filtruj globalną listę `uploaded_files`
        for file_id, file_info in uploaded_files.items():
            if file_info.get('user_id') == session_id:
                user_files_list.append({
                    'fileId': file_id,
                    'originalName': file_info.get('original_name', 'Unknown'),
                    'mimeType': file_info.get('mime_type', 'application/octet-stream'),
                    'category': file_info.get('category', 'general'),
                    'size': file_info.get('size'),
                    'added': datetime.fromtimestamp(file_info.get('added')).isoformat() if file_info.get('added') else None, # Konwertuj timestamp na ISO string
                    'threadId': file_info.get('thread_id') # Opcjonalnie
                })

        logger.info(f"Znaleziono {len(user_files_list)} plików dla użytkownika {session_id}")
        return jsonify({"files": user_files_list})

    except Exception as e:
        logger.exception(f"Błąd podczas listowania plików użytkownika {session_id}: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/web-search', methods=['POST'])
@require_login
def web_search_endpoint():
    """Wykonuje wyszukiwanie w sieci przez Google API."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']

    try:
        data = request.json
        query = data.get('query', '')
        logger.info(f"Żądanie wyszukiwania Google '{query}' przez {session_id}")

        if not query:
            return jsonify({"error": "No search query provided"}), 400

        # Użyj prawdziwego Google Search API
        search_results = google_search(query, num_results=8)
        
        if not search_results:
            # Fallback do symulowanych wyników, jeśli API nie działa
            logger.warning("Google API nie zwróciło wyników, używam symulacji")
            search_results = [
                {
                    "title": f"Example Result 1 for '{query}'",
                    "url": f"https://example.com/search?q={requests.utils.quote(query)}&result=1",
                    "snippet": f"This is a simulated snippet for '{query}' - Google API may not be configured."
                }
            ]

        logger.info(f"Zwracam {len(search_results)} wyników wyszukiwania dla '{query}'")
        return jsonify({"results": search_results})

    except Exception as e:
        logger.exception(f"Błąd podczas wyszukiwania dla {session_id}: {e}")
        return jsonify({"error": str(e)}), 500
        
@app.route('/scrape-url', methods=['POST'])
@require_login
def scrape_url_endpoint():
    """Pobiera i analizuje treść strony internetowej."""
    if client is None:
        logger.error("OpenAI client nie jest zainicjalizowany")
        return jsonify({"error": "Service temporarily unavailable"}), 503
    
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    thread_id = session.get('thread_id') # Pobierz wątek do potencjalnej analizy AI

    if not is_valid_thread_id(thread_id) or not _OPENAI_CLIENT.thread_check(thread_id):
        logger.warning(f"Nieprawidłowy lub nieistniejący thread_id: {thread_id}. Tworzę nowy.")
        thread_id = create_new_thread()
        if not thread_id:
            return jsonify({"error": "AI Assistant currently unavailable"}), 503
        session['thread_id'] = thread_id
        session.modified = True
        logger.info(f"Utworzono nowy thread_id dla scrape_url: {thread_id}")

    try:
        data = request.json
        url = data.get('url')
        category = data.get('category', '')
        logger.info(f"Żądanie scrapowania URL '{url}' przez {session_id} dla kat. '{category}' (wątek: {thread_id})")


        if not url:
            return jsonify({"success": False, "error": "No URL provided"}), 400

        # Prosta walidacja URL - czy zaczyna się od http/https
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url # Spróbuj dodać https
            logger.info(f"Dodano 'https://' do URL: {url}")

        # Pobierz treść strony
        try:
            headers = { # Udawaj przeglądarkę
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.127 Safari/537.36'
            }
            response = requests.get(url, headers=headers, timeout=15, allow_redirects=True) # Zwiększony timeout i śledzenie przekierowań
            response.raise_for_status()  # Sprawdzenie czy nie ma błędu HTTP (4xx, 5xx)

            # Sprawdź content type - czy to HTML?
            content_type = response.headers.get('Content-Type', '').lower()
            if 'html' not in content_type:
                logger.warning(f"URL '{url}' zwrócił typ zawartości '{content_type}', oczekiwano HTML.")
                return jsonify({"success": False, "error": f"URL does not seem to be an HTML page (Content-Type: {content_type})"}), 400

            # Parsuj treść HTML za pomocą BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser') # Użyj response.content dla poprawnego dekodowania

            # Usuń niepotrzebne tagi przed ekstrakcją tekstu
            for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'form', 'button', 'iframe', 'noscript']):
                tag.decompose()

            # Wyodrębnij tytuł
            title = soup.title.string.strip() if soup.title else "No title found"

            # Wyodrębnij opis meta
            meta_description = ""
            meta_tag = soup.find('meta', attrs={'name': 'description'})
            if meta_tag and meta_tag.get('content'):
                meta_description = meta_tag.get('content').strip()

            # Konwertuj główną treść HTML do czystego tekstu
            # Spróbuj znaleźć główny kontener treści (często <main>, <article>, lub div z odpowiednim id/klasą)
            main_content_tag = soup.find('main') or soup.find('article') or soup.find('div', id='content') or soup.find('div', class_='content') or soup.body
            if not main_content_tag: main_content_tag = soup # Fallback do całego soup

            converter = html2text.HTML2Text()
            converter.ignore_links = True # Ignoruj linki w treści tekstowej
            converter.ignore_images = True
            converter.body_width = 0 # Nie zawijaj wierszy
            text_content = converter.handle(str(main_content_tag)).strip()

            # Usuń nadmiarowe puste linie
            text_content = "\n".join([line for line in text_content.splitlines() if line.strip()])


            # Skróć tekst, jeśli jest zbyt długi (dla podglądu i promptu AI)
            preview_max_length = 1000
            ai_prompt_max_length = 8000 # Więcej dla AI
            text_content_preview = text_content[:preview_max_length] + ('...' if len(text_content) > preview_max_length else '')
            text_content_for_ai = text_content[:ai_prompt_max_length] + ('...' if len(text_content) > ai_prompt_max_length else '')

            # Wykonaj analizę AI, jeśli jest wątek
            ai_analysis = ""
            if thread_id:
                try:
                    logger.info(f"Rozpoczynam analizę AI dla URL {url} w wątku {thread_id}")
                    # Przygotuj kontekst dla AI
                    analysis_prompt = (
                        f"Analyze the following website content scraped from {url} "
                        f"(Title: '{title}', Description: '{meta_description}') "
                        f"specifically for the '{category if category else 'general pitch deck'}' section. "
                        f"Provide a concise summary and highlight key points relevant to this section.\n\n"
                        f"Content preview:\n{text_content_for_ai}"
                    )

                    # Wyślij wiadomość do AI
                    client.beta.threads.messages.create(
                        thread_id=thread_id, role="user", content=analysis_prompt
                    )
                    # Uruchom asystenta
                    run = client.beta.threads.runs.create(thread_id=thread_id, assistant_id=assistant_id)
                    logger.info(f"Uruchomiono AI Run {run.id} dla analizy URL.")

                    # Monitoruj status (jak w /ask i /process-files)
                    run_status = "queued"
                    max_attempts = 120; attempts = 0
                    while run_status in ["queued", "in_progress"] and attempts < max_attempts:
                        time.sleep(0.5)
                        run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
                        run_status = run_info.status; attempts += 1
                        if attempts % 10 == 0: logger.debug(f"AI URL Analysis status {run.id} (try {attempts}): {run_status}")

                    if run_status == "completed":
                        messages = client.beta.threads.messages.list(thread_id=thread_id, order="desc")
                        assistant_messages = [m for m in messages.data if m.role == "assistant" and m.run_id == run.id]
                        if assistant_messages and assistant_messages[0].content and hasattr(assistant_messages[0].content[0], 'text'):
                            ai_analysis = assistant_messages[0].content[0].text.value
                            logger.info(f"Otrzymano analizę AI dla URL {url}")
                        else: logger.warning(f"Brak odpowiedzi AI po udanej analizie URL {run.id}")
                    elif run_status == "failed":
                        run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
                        error_message = run_info.last_error.message if run_info.last_error else "Unknown AI error"
                        logger.error(f"AI URL Analysis {run.id} failed: {error_message}")
                    else: logger.error(f"AI URL Analysis {run.id} timed out or ended in status: {run_status}")

                except Exception as e:
                    logger.exception(f"Błąd podczas analizy URL przez OpenAI: {e}")
                    ai_analysis = "" # Kontynuuj bez analizy AI w razie błędu

            logger.info(f"Zakończono scrapowanie URL: {url}. Tytuł: {title}")
            return jsonify({
                "success": True,
                "title": title,
                "description": meta_description,
                "content": text_content_preview, # Zwróć podgląd do UI
                "url": url,
                "category": category,
                "ai_analysis": ai_analysis # Może być pusty string
            })

        except requests.exceptions.Timeout:
            logger.error(f"Timeout fetching URL {url}")
            return jsonify({"success": False, "error": f"Timeout trying to reach the URL: {url}"}), 408 # Request Timeout
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching URL {url}: {e}")
            return jsonify({"success": False, "error": f"Failed to fetch or access URL: {str(e)}"}), 400
        except Exception as e: # Złap inne błędy parsowania itp.
             logger.exception(f"Error processing URL {url}: {e}")
             return jsonify({"success": False, "error": f"Error processing page content: {str(e)}"}), 500

    except Exception as e:
        logger.exception(f"General error in scrape_url_endpoint: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

# --- Nowy endpoint do obsługi zapytań file-qa ---
@app.route('/file-qa', methods=['POST'])
@require_login
def file_qa_endpoint():
    """Obsługuje zapytania do asystenta AI dotyczące zawartości pliku."""
    if client is None:
        logger.error("OpenAI client nie jest zainicjalizowany")
        return jsonify({"error": "Service temporarily unavailable"}), 503
    
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        # Sprawdź, czy plik został przesłany
        if 'file' not in request.files:
            logger.warning(f"Brak pliku w zapytaniu od {session_id}")
            return jsonify({"success": False, "error": "No file part"}), 400
            
        file = request.files['file']
        if file.filename == '':
            logger.warning(f"Przesłano plik bez nazwy od {session_id}")
            return jsonify({"success": False, "error": "No selected file"}), 400
            
        # Sprawdź typ pliku
        allowed_extensions = {'.pdf', '.docx', '.txt'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            logger.warning(f"Niewspierany typ pliku: {file_ext} od {session_id}")
            return jsonify({"success": False, "error": "Unsupported file type. Please upload PDF, DOCX or TXT file."}), 400
            
        # Pobierz pytanie użytkownika
        question = request.form.get('question')
        if not question:
            logger.warning(f"Brak pytania w zapytaniu od {session_id}")
            return jsonify({"success": False, "error": "No question provided"}), 400
        
        # Opcjonalnie użyj istniejącego wątku
        thread_id = request.form.get('thread_id') or session.get('thread_id')
        
        # Tworzymy unikalny identyfikator pliku
        file_id = str(uuid.uuid4())
        
        # Tworzymy katalog dla użytkownika, jeśli nie istnieje
        user_folder = os.path.join(UPLOAD_FOLDER, session_id)
        if not os.path.exists(user_folder):
            os.makedirs(user_folder)
            
        # Ścieżka do pliku
        safe_filename = file_id + file_ext
        file_path = os.path.join(user_folder, safe_filename)
        
        # Zapisujemy plik tymczasowo
        file.save(file_path)
        logger.info(f"Zapisano plik jako: {file_path}")
        
        # Przesłanie pliku do OpenAI
        openai_file_id = upload_file_to_assistant(file_path)
        
        if not openai_file_id:
            return jsonify({"success": False, "error": "Failed to upload file to OpenAI"}), 500
            
        # Zadanie pytania o zawartość pliku
        result = ask_question_about_file(openai_file_id, question, thread_id)
        
        # Zapisujemy informacje o pliku
        file_info = {
            'original_name': file.filename,
            'path': file_path,
            'user_id': session_id,
            'openai_file_id': openai_file_id,
            'thread_id': result.get('thread_id'),
            'category': 'file-qa',
            'mime_type': file.content_type or mimetypes.guess_type(file.filename)[0] or 'application/octet-stream',
            'size': os.path.getsize(file_path),
            'added': time.time()
        }
        uploaded_files[file_id] = file_info
        
        # Aktualizuj thread_id w sesji, jeśli jest nowy
        if result.get('thread_id') and result.get('thread_id') != thread_id:
            session['thread_id'] = result.get('thread_id')
            session.modified = True
        
        # Jeśli wystąpił błąd
        if 'error' in result:
            return jsonify({
                "success": False, 
                "error": result['error'],
                "fileId": file_id,
                "originalName": file.filename
            }), 500
            
        # Jeśli wszystko poszło dobrze
        return jsonify({
            "success": True,
            "answer": result['answer'],
            "fileId": file_id,
            "originalName": file.filename,
            "threadId": result.get('thread_id')
        })
        
    except Exception as e:
        logger.exception(f"Błąd podczas obsługi file-qa od {session_id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

# --- Endpointy Ustawień (Etap 1) ---

@app.route('/settings/profile/name', methods=['PUT'])
@require_login
def update_profile_name():
    """Aktualizuje nazwę profilu użytkownika z persystencją."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    try:
        data = request.json
        new_name = data.get('name', '').strip()
        if not new_name:
            return jsonify(success=False, error="Profile name cannot be empty."), 400
        if len(new_name) > 100:
             return jsonify(success=False, error="Profile name is too long (max 100 chars)."), 400

        # POPRAWKA: Zapisz w persystentnym storage
        if session_id not in user_profiles:
            user_profiles[session_id] = {}
        user_profiles[session_id]['name'] = new_name
        
        # Również zaktualizuj sesję
        session['user_name'] = new_name
        session.modified = True
        
        logger.info(f"Zaktualizowano nazwę użytkownika {session_id} na '{new_name}' (z persystencją)")
        return jsonify(success=True, new_name=new_name)
    except Exception as e:
        logger.exception(f"Błąd podczas aktualizacji nazwy profilu dla {session_id}: {e}")
        return jsonify(success=False, error="Server error updating profile name."), 500

@app.route('/settings/profile/email/request-change', methods=['POST'])
@require_login
def request_email_change():
    """Rozpoczyna proces zmiany emaila (symulacja)."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    try:
        data = request.json
        new_email = data.get('new_email', '').strip().lower()
        # Prosta walidacja emaila
        if not new_email or '@' not in new_email or '.' not in new_email.split('@')[-1]:
            return jsonify(success=False, error="Invalid email format provided."), 400

        # TODO:
        # 1. Sprawdź, czy email nie jest już używany przez innego użytkownika.
        ## 2. Wygeneruj bezpieczny token.
        # 3. Zapisz token i nowy email powiązany z user_id (np. w tabeli `email_changes`).
        # 4. Wyślij email weryfikacyjny na `new_email` z linkiem zawierającym token.
        #    Użyj biblioteki jak Flask-Mail.
        logger.info(f"Symulacja wysłania linku weryfikacyjnego zmiany emaila na '{new_email}' dla {session_id}")
        # Zwróć sukces - frontend poinformuje użytkownika, aby sprawdził email
        return jsonify(success=True, message=f"Verification link simulation for {new_email}")

    except Exception as e:
        logger.exception(f"Błąd podczas żądania zmiany emaila dla {session_id}: {e}")
        return jsonify(success=False, error="Server error requesting email change."), 500

# Endpoint do faktycznej weryfikacji zmiany emaila (placeholder)
# @app.route('/verify-email-change/<token>', methods=['GET'])
# def verify_email_change(token):
#     # TODO: Zaimplementuj logikę weryfikacji tokenu i aktualizacji emaila
#     return "Email change verification endpoint (Not Implemented)"

@app.route('/settings/preferences/notifications', methods=['PUT'])
@require_login
def update_notification_preference():
    """Aktualizuje preferencje powiadomień email."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    try:
        data = request.json
        is_enabled = data.get('enabled')
        if not isinstance(is_enabled, bool):
            return jsonify(success=False, error="Invalid value for 'enabled' (must be true or false)."), 400

        # Zapisz w sesji
        session['notifications_enabled'] = is_enabled
        session.modified = True
        # TODO: Zapisz w bazie danych
        logger.info(f"Zaktualizowano preferencje powiadomień dla {session_id} na {is_enabled}")
        return jsonify(success=True, notifications_enabled=is_enabled)
    except Exception as e:
        logger.exception(f"Błąd podczas aktualizacji preferencji powiadomień dla {session_id}: {e}")
        return jsonify(success=False, error="Server error updating notification preferences."), 500

@app.route('/settings/preferences/language', methods=['PUT'])
@require_login
def update_language_preference():
    """Aktualizuje preferowany język użytkownika."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    try:
        data = request.json
        lang_code = data.get('language', '').lower()
        if lang_code not in app.config['SUPPORTED_LANGUAGES']:
             logger.warning(f"Nieprawidłowy kod języka '{lang_code}' od {session_id}. Obsługiwane: {app.config['SUPPORTED_LANGUAGES']}")
             return jsonify(success=False, error=f"Unsupported language code. Supported: {', '.join(app.config['SUPPORTED_LANGUAGES'])}."), 400

        # Zapisz w sesji
        session['language'] = lang_code
        session.modified = True
        # TODO: Zapisz w bazie danych
        logger.info(f"Zaktualizowano język dla {session_id} na '{lang_code}'")
        
        # Odświeżenie cache Babel, jeśli go używamy aktywnie
        if 'babel' in locals() and hasattr(babel, 'locale_selector_func') and hasattr(babel.locale_selector_func, 'cache_clear'):
             babel.locale_selector_func.cache_clear()
        
        # Zwróć informację o listę dostępnych języków (dla wygody)
        return jsonify(
            success=True, 
            language=lang_code, 
            available_languages=app.config['SUPPORTED_LANGUAGES']
        )
    except Exception as e:
        logger.exception(f"Błąd podczas aktualizacji języka dla {session_id}: {e}")
        return jsonify(success=False, error="Server error updating language preference."), 500

@app.route('/auth/request-password-reset', methods=['POST'])
# Nie wymaga @require_login, bo użytkownik nie jest zalogowany
def request_password_reset():
    """Rozpoczyna proces resetowania hasła (symulacja)."""
    try:
        data = request.json
        email = data.get('email', '').strip().lower()
        if not email or '@' not in email or '.' not in email.split('@')[-1]:
            return jsonify(success=False, error="Invalid email format provided."), 400

        # TODO:
        # 1. Sprawdź, czy użytkownik z tym emailem istnieje w bazie.
        # 2. Wygeneruj bezpieczny token resetowania hasła.
        # 3. Zapisz token powiązany z user_id i ustaw datę ważności.
        # 4. Wyślij email na adres użytkownika z linkiem zawierającym token
        #    (np. /reset-password/<token>). Użyj Flask-Mail.

        logger.info(f"Symulacja wysłania linku resetowania hasła na adres '{email}'")
        # Zawsze zwracaj sukces, nawet jeśli email nie istnieje, aby nie ujawniać informacji
        return jsonify(success=True, message=f"Password reset link simulation for {email}")

    except Exception as e:
        logger.exception(f"Błąd podczas żądania resetu hasła dla {email if 'email' in locals() else 'unknown'}: {e}")
        # Zwracaj generyczny błąd dla bezpieczeństwa
        return jsonify(success=False, error="Server error processing password reset request."), 500

# --- Endpointy Etapu 2: Zarządzanie Danymi i Prywatnością ---

@app.route('/settings/data/export', methods=['GET'])
@require_login
def export_account_data():
    """Eksportuje dane konta użytkownika do pliku ZIP zgodnie z wymogami RODO/GDPR."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        logger.info(f"Żądanie eksportu danych konta dla użytkownika {session_id}")
        
        # Tworzenie pliku ZIP w pamięci
        mem_zip = BytesIO()
        with zipfile.ZipFile(mem_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
            # 1. Dane profilu użytkownika
            profile_data = {
                "user_id": user_data['user_id'],
                "name": user_data['name'],
                "email": user_data['email'],
                "plan": user_data['plan'],
                "language": user_data['language'],
                "notifications_enabled": user_data['notificationsEnabled'],
                "account_created": "N/A",  # Placeholder - w rzeczywistym systemie pobierz z bazy danych
                "last_login": datetime.now().isoformat()  # Placeholder - w rzeczywistym systemie pobierz z bazy danych
            }
            zf.writestr('profile/user_info.json', json.dumps(profile_data, indent=4))
            
            # 2. Historia konwersacji
            if 'conversation_history' in session:
                conversation_data = session['conversation_history']
                zf.writestr('conversations/history.json', json.dumps(conversation_data, indent=4))
            else:
                zf.writestr('conversations/history.json', json.dumps({}, indent=4))
            
            # 3. Pliki użytkownika (metadane)
            user_files_metadata = []
            for file_id, file_info in uploaded_files.items():
                if file_info.get('user_id') == session_id:
                    user_files_metadata.append({
                        'file_id': file_id,
                        'original_name': file_info.get('original_name'),
                        'mime_type': file_info.get('mime_type'),
                        'size': file_info.get('size'),
                        'uploaded_at': datetime.fromtimestamp(file_info.get('added', 0)).isoformat(),
                        'category': file_info.get('category')
                    })
            
            zf.writestr('files/metadata.json', json.dumps(user_files_metadata, indent=4))
            
            # 4. Faktyczna zawartość plików (jeśli rozmiar nie jest za duży)
            user_folder = os.path.join(UPLOAD_FOLDER, session_id)
            files_added = 0
            total_size = 0
            size_limit = 50 * 1024 * 1024  # Limit 50MB dla całego eksportu
            
            for file_id, file_info in uploaded_files.items():
                if file_info.get('user_id') == session_id:
                    file_path = file_info.get('path')
                    if file_path and os.path.exists(file_path):
                        file_size = os.path.getsize(file_path)
                        
                        # Sprawdź, czy dodanie pliku nie przekroczy limitu
                        if total_size + file_size <= size_limit:
                            original_name = file_info.get('original_name', 'unknown_file')
                            safe_filename = f"{file_id}_{original_name}"
                            zf.write(file_path, f'files/content/{safe_filename}')
                            files_added += 1
                            total_size += file_size
                        else:
                            # Przekroczono limit, dodaj notatkę
                            logger.warning(f"Plik {file_id} pominięty w eksporcie ze względu na limit rozmiaru")
            
            # Dodaj README z informacjami o eksporcie
            readme_content = f"""
            # Eksport danych konta Deckster AI
            
            Data eksportu: {datetime.now().isoformat()}
            Użytkownik: {user_data['name']} ({user_data['email']})
            
            ## Zawartość archiwum:
            
            1. /profile/ - Dane profilu użytkownika
            2. /conversations/ - Historia konwersacji
            3. /files/ - Metadane i zawartość plików
            
            Łączna liczba plików: {files_added}
            
            Ten plik został wygenerowany automatycznie jako część prawa do przenoszenia danych zgodnie z RODO/GDPR.
            """
            
            zf.writestr('README.txt', readme_content.strip())
        
        # Przesuń wskaźnik na początek pliku w pamięci
        mem_zip.seek(0)
        
        # Przygotuj nazwę pliku
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"deckster_data_export_{session_id}_{timestamp}.zip"
        
        logger.info(f"Pomyślnie wygenerowano plik eksportu danych dla użytkownika {session_id}")
        
        # Wyślij plik
        return send_file(
            mem_zip,
            mimetype='application/zip',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.exception(f"Błąd podczas eksportu danych konta dla użytkownika {session_id}: {e}")
        return jsonify({"error": f"Failed to export account data: {str(e)}"}), 500

@app.route('/settings/account/delete', methods=['POST', 'DELETE'])
@require_login
def delete_account():
    """Usuwa konto użytkownika i wszystkie powiązane dane."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        # Weryfikacja tożsamości - wymaga hasła (lub innej metody potwierdzenia)
        data = request.json
        password = data.get('password')
        confirmation_code = data.get('confirmation_code')
        
        if not password and not confirmation_code:
            logger.warning(f"Próba usunięcia konta bez potwierdzenia tożsamości dla {session_id}")
            return jsonify(success=False, error="Password or confirmation code is required to delete your account."), 400
        
        # W rzeczywistej implementacji należy zweryfikować hasło użytkownika
        # Tutaj zakładamy, że weryfikacja przeszła pomyślnie
        
        logger.info(f"Rozpoczęto proces usuwania konta dla użytkownika {session_id}")
        
        # 1. Usuwanie plików użytkownika
        deleted_files = []
        user_folder = os.path.join(UPLOAD_FOLDER, session_id)
        
        # Usuń fizyczne pliki
        if os.path.exists(user_folder):
            for file_id, file_info in list(uploaded_files.items()):
                if file_info.get('user_id') == session_id:
                    file_path = file_info.get('path')
                    if file_path and os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                            deleted_files.append(file_id)
                            logger.debug(f"Usunięto plik: {file_path}")
                        except OSError as e:
                            logger.error(f"Nie udało się usunąć pliku {file_path}: {e}")
            
            # Usuń katalog użytkownika (jeśli pusty)
            try:
                os.rmdir(user_folder)
                logger.debug(f"Usunięto katalog użytkownika: {user_folder}")
            except OSError as e:
                logger.warning(f"Nie udało się usunąć katalogu użytkownika {user_folder}: {e}")
        
        # Usuń metadane plików
        for file_id in deleted_files:
            uploaded_files.pop(file_id, None)
        
        # 2. Usuń wątki konwersacji z OpenAI (jeśli to możliwe)
        thread_id = session.get('thread_id')
        if thread_id:
            try:
                # Próba usunięcia wątku w OpenAI
                # Uwaga: API OpenAI może nie wspierać bezpośredniego usuwania wątków
                # client.beta.threads.delete(thread_id=thread_id)
                logger.info(f"Usunięto wątek OpenAI: {thread_id} (symulacja)")
            except Exception as e:
                logger.error(f"Nie udało się usunąć wątku OpenAI {thread_id}: {e}")
                # Nie przerywaj procesu usuwania konta, jeśli to się nie powiedzie
        
        # 3. Usuń historię konwersacji z sesji
        session.pop('conversation_history', None)
        
        # 4. W rzeczywistej implementacji należałoby usunąć dane użytkownika z bazy danych
        # DELETE FROM users WHERE user_id = ?
        # DELETE FROM user_settings WHERE user_id = ?
        # itd.
        
        # 5. Wyczyść sesję i wyloguj użytkownika
        session.clear()
        
        logger.info(f"Pomyślnie usunięto konto użytkownika {session_id}")
        
        return jsonify(success=True, message="Your account has been successfully deleted.")
        
    except Exception as e:
        logger.exception(f"Błąd podczas usuwania konta dla użytkownika {session_id}: {e}")
        return jsonify(success=False, error="Failed to delete account. Please try again or contact support."), 500

# --- Nowy endpoint do obsługi plików upuszczonych na pole tekstowe ---
@app.route('/process-dropped-file', methods=['POST'])
@require_login
def process_dropped_file():
    """Przetwarza plik upuszczony na pole tekstowe."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    thread_id = session.get('thread_id', 'unknown_thread')
    
    try:
        # Sprawdź czy otrzymaliśmy plik w żądaniu
        if 'file' not in request.files:
            logger.warning(f"Brak pliku w żądaniu od {session_id}")
            return jsonify({"success": False, "error": "No file provided"}), 400
            
        file = request.files['file']
        if file.filename == '':
            logger.warning(f"Przesłano plik bez nazwy od {session_id}")
            return jsonify({"success": False, "error": "Empty filename"}), 400
            
        category = request.form.get('category', 'general')
        logger.info(f"Otrzymano plik upuszczony na pole tekstowe: {file.filename} od {session_id}, kategoria: {category}")
        
        # Generujemy unikalny identyfikator dla pliku
        file_id = str(uuid.uuid4())
        
        # Sprawdzamy czy istnieje folder dla użytkownika
        user_folder = os.path.join(UPLOAD_FOLDER, session_id)
        if not os.path.exists(user_folder):
            os.makedirs(user_folder)
            
        # Zapisujemy plik
        file_extension = os.path.splitext(file.filename)[1]
        safe_filename = file_id + file_extension
        file_path = os.path.join(user_folder, safe_filename)
        file.save(file_path)
        
        # Zapisujemy informacje o pliku
        file_info = {
            'id': file_id,
            'original_name': file.filename,
            'path': file_path,
            'user_id': session_id,
            'thread_id': thread_id,
            'category': category,
            'mime_type': file.content_type or mimetypes.guess_type(file.filename)[0] or 'application/octet-stream',
            'size': os.path.getsize(file_path),
            'added': time.time(),
            'name': file.filename
        }
        uploaded_files[file_id] = file_info
        logger.info(f"Zapisano metadane pliku {file_id}: {file_info}")
        
        # Odczytaj zawartość pliku w zależności od typu
        file_data = {}
        try:
            if file_info['mime_type'].startswith('text/'):
                # Dla plików tekstowych
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_data['content'] = f.read()
            elif file_info['mime_type'].startswith('image/'):
                # Dla obrazów - konwersja do base64
                with open(file_path, 'rb') as f:
                    file_content = f.read()
                    file_data['content'] = base64.b64encode(file_content).decode('utf-8')
                    file_data['contentType'] = file_info['mime_type']
            # Możemy dodać więcej typów plików w przyszłości
        except Exception as e:
            logger.warning(f"Nie udało się odczytać zawartości pliku {file_id}: {e}")
            # Nie zwracaj błędu, kontynuuj bez zawartości
            
        # Dodaj plik do listy plików użytkownika dla widoku panelu plików
        userFiles = []
        userFiles.append({
            'id': file_id,
            'name': file.filename,
            'size': file_info['size'],
            'type': file_info['mime_type'],
            'added': datetime.fromtimestamp(file_info['added']).isoformat(),
            'category': category
        })
        
        return jsonify({
            "success": True,
            "fileId": file_id,
            "originalName": file.filename,
            "fileData": file_data,
            "fileInfo": file_info
        })
        
    except Exception as e:
        logger.exception(f"Błąd podczas przetwarzania upuszczonego pliku od {session_id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500
        
# --- Endpoint do obsługi ogólnego chatu ---
@app.route('/general-chat', methods=['POST'])
@require_login
def general_chat():
    """Obsługuje ogólny chat z poprawną obsługą thread_id."""
    if client is None:
        return jsonify({"error": "Service temporarily unavailable"}), 503
    
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        data = request.json
        message = data.get('message')
        request_thread_id = data.get('threadId')
        
        if not message:
            return jsonify({"error": "No message provided"}), 400
            
        # UJEDNOLICONA OBSŁUGA THREAD_ID
        current_thread_id = request_thread_id or session.get('thread_id')
        
        try:
            thread_id, is_new = ensure_valid_thread(session_id, current_thread_id)
            
            if thread_id != session.get('thread_id'):
                session['thread_id'] = thread_id
                session.modified = True
                
        except RuntimeError as e:
            return jsonify({"error": "AI Assistant currently unavailable"}), 503
        
        # Wyślij do OpenAI
        client.beta.threads.messages.create(
            thread_id=thread_id,
            role="user",
            content=message
        )
        
        run = client.beta.threads.runs.create(
            thread_id=thread_id,
            assistant_id=assistant_id
        )
        
        # Monitoruj (jak wyżej)
        run_status = "queued"
        max_attempts = 120
        attempts = 0
        
        while run_status in ["queued", "in_progress"] and attempts < max_attempts:
            time.sleep(0.5)
            run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
            run_status = run_info.status
            attempts += 1
            
        if run_status == "completed":
            messages = client.beta.threads.messages.list(thread_id=thread_id, order="desc")
            assistant_messages = [msg for msg in messages.data if msg.role == "assistant"]
            
            if assistant_messages and assistant_messages[0].content:
                answer = assistant_messages[0].content[0].text.value
                
                # Zapisz w historii ogólnego chatu
                if session_id not in general_chat_history:
                    general_chat_history[session_id] = []
                    
                general_chat_history[session_id].append({
                    "question": message,
                    "answer": answer,
                    "timestamp": time.time()
                })
                
                # ZAWSZE zwróć threadId
                return jsonify({
                    "answer": answer, 
                    "threadId": thread_id
                })
            else:
                return jsonify({"error": "No response from assistant"}), 500
        else:
            error_message = "Assistant failed to respond in time"
            if run_status == "failed":
                run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
                if run_info.last_error:
                    error_message = run_info.last_error.message
                    
            return jsonify({"error": error_message}), 500
            
    except Exception as e:
        logger.exception(f"Error in general_chat for {session_id}: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500
        
# --- Endpoint do pobierania historii ogólnego chatu ---
@app.route('/get-general-chat-history', methods=['GET'])
@require_login
def get_general_chat_history():
    """Pobiera historię ogólnego chatu dla zalogowanego użytkownika."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        # Pobierz historię ogólnego chatu dla użytkownika
        history = general_chat_history.get(session_id, [])
        
        return jsonify({
            "success": True,
            "history": history
        })
        
    except Exception as e:
        logger.exception(f"Błąd podczas pobierania historii ogólnego chatu dla {session_id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

# --- Nowy endpoint do analizy załączników w rozmowie ---
@app.route('/analyze-conversation-files', methods=['POST'])
@require_login
def analyze_conversation_files():
    """Analizuje załączniki w kontekście rozmowy."""
    if client is None:
        logger.error("OpenAI client nie jest zainicjalizowany")
        return jsonify({"error": "Service temporarily unavailable"}), 503
    
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    thread_id = session.get('thread_id')
    
    try:
        data = request.json
        file_ids = data.get('fileIds', [])
        context = data.get('context', '')
        
        if not file_ids:
            return jsonify({"success": False, "error": "No files specified for analysis"}), 400
        
        if not is_valid_thread_id(thread_id) or not _OPENAI_CLIENT.thread_check(thread_id):
            logger.warning(f"Nieprawidłowy lub nieistniejący thread_id: {thread_id}. Tworzę nowy.")
            thread_id = create_new_thread()
            if not thread_id:
                logger.error(f"Nie udało się utworzyć wątku dla {session_id}")
                return jsonify({"error": "AI Assistant currently unavailable. Please try again later."}), 503
            
            session['thread_id'] = thread_id
            session.modified = True
            logger.info(f"Zapisano nowy thread_id {thread_id} w sesji dla {session_id}")
            
        logger.info(f"Żądanie analizy {len(file_ids)} plików w kontekście rozmowy od {session_id}")
        
        # Przygotuj informacje o plikach
        valid_files = []
        file_contents = []
        
        for file_id in file_ids:
            if file_id in uploaded_files:
                file_info = uploaded_files[file_id]
                
                # Sprawdź, czy plik należy do użytkownika
                if file_info.get('user_id') == session_id:
                    valid_files.append({
                        'id': file_id,
                        'name': file_info.get('original_name', 'Unknown'),
                        'type': file_info.get('mime_type', 'application/octet-stream')
                    })
                    
                    # Przygotuj zawartość pliku, jeśli możliwe
                    file_path = file_info.get('path')
                    try:
                        if file_path and os.path.exists(file_path):
                            mime_type = file_info.get('mime_type', 'application/octet-stream')
                            
                            if mime_type.startswith('text/'):
                                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                    content = f.read()
                                    file_contents.append({
                                        'name': file_info.get('original_name', 'Unknown'),
                                        'type': mime_type,
                                        'content': content
                                    })
                            elif mime_type.startswith('image/'):
                                file_contents.append({
                                    'name': file_info.get('original_name', 'Unknown'),
                                    'type': mime_type,
                                    'content': 'This is an image file. Please analyze what you see in it.'
                                })
                            elif mime_type.startswith('application/pdf'):
                                file_contents.append({
                                    'name': file_info.get('original_name', 'Unknown'),
                                    'type': mime_type,
                                    'content': 'This is a PDF document. Please extract and analyze its content.'
                                })
                            else:
                                # Dla innych typów plików
                                file_contents.append({
                                    'name': file_info.get('original_name', 'Unknown'),
                                    'type': mime_type,
                                    'content': f'This is a {mime_type} file. Please analyze its structure and content if possible.'
                                })
                    except Exception as e:
                        logger.warning(f"Nie udało się odczytać zawartości pliku {file_id}: {e}")
                        # Kontynuuj bez zawartości tego pliku
                else:
                    logger.warning(f"Użytkownik {session_id} próbuje analizować plik {file_id}, do którego nie ma dostępu.")
            else:
                logger.warning(f"Plik o ID {file_id} nie znaleziony w metadanych.")
                
        if not valid_files:
            return jsonify({"success": False, "error": "No valid files found for analysis"}), 400
            
        # Przygotuj wiadomość dla asystenta
        analysis_prompt = f"Please analyze the following files in the context of our conversation.\n\n"
        
        if context:
            analysis_prompt += f"User says: {context}\n\n"
            
        analysis_prompt += "Files to analyze:\n"
        for i, file in enumerate(valid_files, 1):
            analysis_prompt += f"{i}. {file['name']} ({file['type']})\n"
            
        # Dodaj zawartość plików
        if file_contents:
            analysis_prompt += "\n### File Contents:\n\n"
            for idx, file_content in enumerate(file_contents, 1):
                analysis_prompt += f"#### File {idx}: {file_content['name']} ({file_content['type']})\n\n"
                analysis_prompt += file_content['content']
                analysis_prompt += "\n\n---\n\n"
                
        # Wyślij wiadomość do asystenta
        client.beta.threads.messages.create(
            thread_id=thread_id,
            role="user",
            content=analysis_prompt
        )
        
        # Uruchom asystenta
        run = client.beta.threads.runs.create(
            thread_id=thread_id,
            assistant_id=assistant_id
        )
        
        # Monitoruj status
        run_status = "queued"
        max_attempts = 120
        attempts = 0
        
        while run_status in ["queued", "in_progress"] and attempts < max_attempts:
            time.sleep(0.5)
            run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
            run_status = run_info.status
            attempts += 1
            if attempts % 10 == 0:
                logger.debug(f"Sprawdzanie statusu analizy plików {run.id} (próba {attempts}): {run_status}")
                
        if run_status == "completed":
            messages = client.beta.threads.messages.list(thread_id=thread_id, order="desc")
            assistant_messages = [msg for msg in messages.data if msg.role == "assistant"]
            
            if assistant_messages:
                answer = assistant_messages[0].content[0].text.value
                logger.info(f"Otrzymano odpowiedź AI dla analizy plików")
                
                # Zapisz historię zapytania i odpowiedzi
                if session_id not in general_chat_history:
                    general_chat_history[session_id] = []
                    
                general_chat_history[session_id].append({
                    "question": f"Analyze the following files: {', '.join(f['name'] for f in valid_files)}",
                    "answer": answer,
                    "timestamp": time.time(),
                    "files_analyzed": file_ids
                })
                
                return jsonify({
                    "success": True,
                    "analysis": answer,
                    "files_analyzed": len(valid_files)
                })
            else:
                return jsonify({"success": False, "error": "Assistant did not provide an analysis"}), 500
        elif run_status == "failed":
            run_info = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
            error_message = run_info.last_error.message if run_info.last_error else "Unknown error"
            logger.error(f"File analysis run {run.id} failed: {error_message}")
            return jsonify({"success": False, "error": f"Analysis failed: {error_message}"}), 500
        else:
            logger.error(f"File analysis timed out or ended in unexpected status: {run_status}")
            return jsonify({"success": False, "error": f"Analysis timed out (status: {run_status})"}), 504
    
    except Exception as e:
        logger.exception(f"Błąd podczas analizy plików dla {session_id}: {e}")
        return jsonify({"success": False, "error": f"Server error: {str(e)}"}), 500

@app.route('/file-analysis-modal', methods=['GET'])
@require_login
def file_analysis_modal():
    """Dostarcza listę plików użytkownika dla modalu analizy."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        user_files_list = []
        
        # Filtruj globalną listę uploaded_files dla użytkownika
        for file_id, file_info in uploaded_files.items():
            if file_info.get('user_id') == session_id:
                user_files_list.append({
                    'fileId': file_id,
                    'originalName': file_info.get('original_name', 'Unknown'),
                    'mimeType': file_info.get('mime_type', 'application/octet-stream'),
                    'size': file_info.get('size', 0),
                    'added': datetime.fromtimestamp(file_info.get('added', 0)).isoformat()
                })
                
        # Sortuj pliki według daty dodania (najnowsze pierwsze)
        user_files_list.sort(key=lambda x: x['added'], reverse=True)
        
        return jsonify({
            "success": True,
            "files": user_files_list
        })
        
    except Exception as e:
        logger.exception(f"Błąd podczas pobierania plików dla modalu analizy: {e}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

# --- Nowy endpoint do pobierania historii konwersacji ---
@app.route('/get-conversation-history', methods=['GET'])
@require_login
def get_conversation_history():
    """Pobiera historię konwersacji dla zalogowanego użytkownika."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        # Historię przechowujemy w sesji użytkownika
        conversation_history = session.get('conversation_history', {})
        
        return jsonify({
            "success": True,
            "history": conversation_history
        })
        
    except Exception as e:
        logger.exception(f"Błąd podczas pobierania historii konwersacji dla {session_id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500
        
# --- Nowy endpoint do zapisywania konwersacji ---
@app.route('/save-conversation', methods=['POST'])
@require_login
def save_conversation():
    """Zapisuje konwersację dla określonej kategorii."""
    user_data = get_current_user_data()
    session_id = user_data['user_id']
    
    try:
        data = request.json
        category = data.get('category')
        content = data.get('content')
        
        if not category or not content:
            return jsonify({"success": False, "error": "Missing category or content"}), 400
            
        # Zapisujemy historię konwersacji w sesji
        if 'conversation_history' not in session:
            session['conversation_history'] = {}
            
        session['conversation_history'][category] = {
            'content': content,
            'timestamp': time.time()
        }
        
        # Oznaczamy sesję jako zmodyfikowaną, aby zmiany zostały zapisane
        session.modified = True
        
        logger.info(f"Zapisano konwersację dla kategorii: {category}")
        return jsonify({"success": True})
        
    except Exception as e:
        logger.exception(f"Błąd podczas zapisywania konwersacji dla {session_id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

# --- Główny Punkt Uruchomienia ---
if __name__ == '__main__':
    # Użyj Gunicorn lub innego serwera WSGI w produkcji!
    # Ustawienia debug=True i host='0.0.0.0' są dla dewelopmentu.
    port = int(os.environ.get('PORT', 5000)) # Użyj portu z env lub domyślnie 5000
    logger.info(f"Uruchamianie serwera Flask na porcie {port}...")
    
    app.run(debug=False, host='0.0.0.0', port=port)
